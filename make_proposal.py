from docx import Document
from docx.shared import Pt

proposal_text = {
    'title': 'Sample-Efficient FunSearch for CVRP',
    'author': '秦子桁 - 123',
    'content': [
        ("Project and Motivation", "We propose to adapt FunSearch — an LLM-driven program search method for automatic heuristic design — to the Capacitated Vehicle Routing Problem (CVRP), with a focus on sample-efficient evaluation. CVRP is a classical combinatorial optimization problem with widely-used benchmarks (e.g., CVRPLib, Uchoa instances). Improving FunSearch's sample efficiency reduces computational cost and enables broader application to routing problems."),
        ("Which topic and why", "We choose 'sample-efficient FunSearch for CVRP' because CVRP offers rich, well-studied benchmarks and practical relevance. The project builds on FunSearch while contributing a concrete efficiency improvement (duplicate-detection and smarter sampling) that has both theoretical and empirical value."),
        ("Tentative Plan", "1) Implement a FunSearch pipeline for CVRP: template heuristic wrapper + LLM interface for program generation/variation. 2) Design a lightweight duplicate/functional-similarity detector to avoid re-evaluating functionally equivalent programs. 3) Develop sampling and evaluation strategies (adaptive sampling, early stopping) to reduce sandbox runs. 4) Integrate baseline solvers (OR-Tools, LKH/HGS) for comparison and run controlled experiments on benchmark instances."),
        ("Evaluation Plan", "Datasets: CVRPLib (Uchoa sets A,B,E,F,M,P recommended). Metrics: total route distance, number of vehicles, runtime, and gap to best-known solutions. Baselines: Google OR-Tools, LKH (via pylkh), and HGS if available. Experiments: compare average performance vs. baselines, measure reduction in sandbox evaluations, and ablation tests for duplicate-detection and sampling strategies."),
        ("Deliverables & Timeline", "Deliver a reproducible codebase (Colab/GitHub), experiments on selected Uchoa instances, and a short report. Milestones: Week 1–2: environment, dataset ingestion, baseline runner; Week 3–5: FunSearch implementation + duplicate detector; Week 6–8: experiments, ablations, and write-up."),
        ("Contact", "秦子桁 — 123")
    ]
}

# Create document
doc = Document()

# Title
p = doc.add_paragraph()
run = p.add_run(proposal_text['title'])
run.bold = True
run.font.size = Pt(14)

# Author
p = doc.add_paragraph()
run = p.add_run(proposal_text['author'])
run.italic = True
run.font.size = Pt(11)

doc.add_paragraph('')

for heading, para in proposal_text['content']:
    h = doc.add_paragraph()
    r = h.add_run(heading + ':')
    r.bold = True
    r.font.size = Pt(11)
    t = doc.add_paragraph(para)
    t.style.font.size = Pt(10)

# Save
out = 'CVRP_FunSearch_Proposal.docx'
doc.save(out)
print('Saved', out)
