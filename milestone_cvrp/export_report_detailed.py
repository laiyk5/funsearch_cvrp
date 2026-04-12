from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """Add heading with proper styling."""
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.size = Pt(13 if level == 1 else 12 if level == 2 else 11)


def add_text(doc: Document, text: str, size: int = 11, bold: bool = False, color: tuple | None = None) -> None:
    """Add paragraph with optional styling."""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor(*color)


def add_table(doc: Document, rows: list[list[str]], header: bool = True) -> None:
    """Add a simple table."""
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    
    for i, row_data in enumerate(rows):
        for j, cell_data in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell_para = cell.paragraphs[0]
            cell_para.text = str(cell_data)
            if header and i == 0:
                for run in cell_para.runs:
                    run.bold = True
                    run.font.size = Pt(10)


def main() -> None:
    out_dir = Path("outputs")
    results_path = out_dir / "milestone_results.json"
    history_path = out_dir / "search_history.json"

    if not results_path.exists():
        raise FileNotFoundError("Run run_milestone.py first")

    results = json.loads(results_path.read_text(encoding="utf-8"))
    history = json.loads(history_path.read_text(encoding="utf-8"))

    baseline = results["baseline"]
    method = results["sample_efficient_funsearch"]
    improvement = results["improvement"]

    doc = Document()

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Milestone Report: Sample-Efficient FunSearch for CVRP")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Group members
    doc.add_paragraph()
    add_heading(doc, "Group Members", level=2)
    for name in ["QIN YUANCHENG (59061061)", "QIN ZIHENG (59866035)", "LAI YIKAI (59563061)"]:
        doc.add_paragraph(name, style="List Bullet")

    # 1. Problem Description and Motivation
    doc.add_paragraph()
    add_heading(doc, "1. Problem Description and Motivation", level=2)

    add_text(doc, "1.1 Problem Definition")
    add_text(
        doc,
        "The Capacitated Vehicle Routing Problem (CVRP) is a fundamental combinatorial optimization problem in operations research and logistics. Formally, given a depot (node 0) and a set of customers {1, 2, ..., n} each with a demand d_i, we seek to determine a set of routes for a homogeneous fleet of vehicles (each with fixed capacity Q) such that: (1) every customer is visited exactly once, (2) vehicle capacity constraints are satisfied, and (3) the total distance traveled is minimized.",
        size=10,
    )

    add_text(doc, "1.2 Motivation for Sample-Efficient Search", bold=True)
    add_text(
        doc,
        "Traditional CVRP solution methods rely on human expertise to design effective heuristics through trial-and-error. Recent work in automated algorithm design (e.g., FunSearch from Google DeepMind) demonstrates that large language models can search over the space of heuristic algorithms and discover novel, high-performing solutions. However, LLM-based search is computationally expensive due to extensive evaluation of candidate heuristics. Our key observation is that many generated heuristics exhibit functionally similar behavior despite syntactic differences, leading to redundant evaluations. This milestone focuses on reducing evaluation cost through: (1) duplicate detection using functional signatures, and (2) two-stage evaluation with early pruning of weak candidates.",
        size=10,
    )

    # 2. Method/Approach Design
    doc.add_paragraph()
    add_heading(doc, "2. Method/Approach Design", level=2)

    add_text(doc, "2.1 Heuristic Family and Parameterization", bold=True)
    add_text(
        doc,
        "We consider a family of greedy constructive heuristics for CVRP parameterized by three weights (w1, w2, w3). The heuristic sequentially builds routes by greedily selecting the next node:\n"
        "  score(c) = w1 * (-distance_current_to_c) + w2 * (demand_c / capacity) + w3 * (-distance_c_to_depot)\n"
        "The weights control the relative importance of: proximity to current node, customer demand urgency, and proximity to returning to depot. This parameterization balances exploitation and exploration.",
        size=10,
    )

    add_text(doc, "2.2 Sample-Efficient Search Algorithm", bold=True)
    add_text(
        doc,
        "Algorithm: Sample-Efficient FunSearch\n"
        "1. Initialize population P with k random weight vectors, evaluate on early subset E ⊂ instances\n"
        "2. evaluated_sigs ← empty set\n"
        "3. For iteration i = 1 to max_iterations:\n"
        "   a. parent ← select_top_k(P)  # pick from best performers\n"
        "   b. w_new ← mutate(parent.weights)  # Gaussian perturbation\n"
        "   c. sig ← round(w_new, 2)  # functional signature\n"
        "   d. If sig ∈ evaluated_sigs, skip (duplicate avoidance)\n"
        "   e. score_early ← evaluate(w_new, E)  # fast evaluation on subset\n"
        "   f. If score_early promising (relative to P_top_k), then:\n"
        "       - score_full ← evaluate(w_new, all instances)  # expensive full eval\n"
        "       - Add candidate to P, prune to maintain diversity\n"
        "   g. Else, record as 'early_pruned'\n"
        "4. Return best candidate and search history",
        size=10,
    )

    add_text(doc, "2.3 Key Innovation: Early Pruning + Duplicate Filtering", bold=True)
    add_text(
        doc,
        "To reduce evaluation cost, we implement two mechanisms:\n"
        "• Early Pruning: Evaluate candidate heuristics on a small subset of instances first. If performance is weak relative to current best, skip expensive full evaluation.\n"
        "• Duplicate Filtering: Use rounded weight vectors as functional signatures. If a signature is already evaluated, skip re-evaluation.\n"
        "Combined, these reduce evaluation count while maintaining solution quality.",
        size=10,
    )

    # 3. Experimental Setup
    doc.add_paragraph()
    add_heading(doc, "3. Experimental Setup", level=2)

    add_text(doc, "3.1 Benchmark Instances", bold=True)
    add_text(
        doc,
        f"For this milestone, we use synthetic CVRP instances for reproducibility: {', '.join(results['dataset'])}. Each instance is generated with a fixed random seed and contains customers with random coordinates and demands. This ensures repeatability and compliance with the course milestone timeline.",
        size=10,
    )

    add_text(doc, "3.2 Baseline Methods", bold=True)
    add_text(
        doc,
        "Baseline (Nearest Neighbor): A classical constructive heuristic that always selects the nearest unvisited feasible customer. This provides a strong and interpretable baseline.",
        size=10,
    )

    add_text(doc, "3.3 Evaluation Metrics", bold=True)
    add_text(
        doc,
        "• Avg Total Distance: average travel distance across all instances.\n"
        "• Avg Number of Vehicles: average number of routes used.\n"
        "• Cost Score: weighted objective = avg_distance + 20 * avg_num_routes (favors fewer vehicles).\n"
        "• Sample Efficiency: number of unique candidates evaluated (lower is better for time/cost).",
        size=10,
    )

    # 4. Preliminary Results
    doc.add_paragraph()
    add_heading(doc, "4. Preliminary Results", level=2)

    add_text(doc, "4.1 Performance Comparison", bold=True)
    
    results_table = [
        ["Method", "Avg Distance", "Avg #Routes", "Cost Score"],
        [baseline["name"], str(baseline["avg_distance"]), str(baseline["avg_num_routes"]), "N/A"],
        ["Sample-Efficient FunSearch", str(method["avg_distance"]), str(method["avg_num_routes"]), str(method["best_score"])],
    ]
    add_table(doc, results_table, header=True)

    add_text(doc, "4.2 Improvement Analysis", bold=True)
    improvements_text = (
        f"• Distance reduction: {improvement['distance_delta']:.3f} units ({improvement['distance_improvement_percent']:.2f}%)\n"
        f"• Vehicle count: no change (both methods use {baseline['avg_num_routes']} routes on average)\n"
        f"• Best weight vector discovered: {[round(w, 4) for w in method['best_weights']]}\n"
        f"• Total unique candidates evaluated: {method['unique_candidates']}\n"
        f"• Early-pruned candidates: see search_history.json for details"
    )
    add_text(doc, improvements_text, size=10)

    add_text(doc, "4.3 Search Behavior", bold=True)
    pruned_count = sum(1 for h in history if h.get("status") == "early_pruned")
    full_eval_count = sum(1 for h in history if h.get("status") == "full_eval")
    add_text(
        doc,
        f"During the search, we evaluated {full_eval_count} candidates fully and pruned {pruned_count} candidates early. This represents a sampling efficiency improvement: we avoided {pruned_count} expensive full evaluations (~{100*pruned_count/(pruned_count+full_eval_count):.1f}% of early-stage candidates).",
        size=10,
    )

    # 5. Technical Contributions
    doc.add_paragraph()
    add_heading(doc, "5. Technical Contributions (Novelty)", level=2)

    add_text(doc, "• Dual-stage evaluation strategy: Unlike naive FunSearch that evaluates every candidate fully, we use cheap early-stage filtering.", size=10)
    add_text(doc, "• Functional signature-based duplicate detection: Rounded weight vectors serve as signatures to avoid redundant evaluations.", size=10)
    add_text(doc, "• Reproducible baseline: Synthetic benchmarks with fixed seeds for open reproducibility.", size=10)

    # 6. Limitations and Future Work
    doc.add_paragraph()
    add_heading(doc, "6. Limitations and Future Work", level=2)

    add_text(doc, "• Synthetic benchmarks: Current milestone uses generated instances. Final project will include CVRPLib/Uchoa standard benchmarks.", size=10)
    add_text(doc, "• Limited baselines: Will integrate Google OR-Tools and (if available) LKH for stronger comparison.", size=10)
    add_text(doc, "• Functional similarity: Rounded signatures are a proxy; future work will explore semantic code similarity metrics.", size=10)
    add_text(doc, "• Scalability: Current implementation tested on small instances (20–36 customers); final submission will scale to larger instances.", size=10)

    # 7. Code Organization
    doc.add_paragraph()
    add_heading(doc, "7. Code Organization and Submission", level=2)

    code_files = [
        ["File", "Purpose"],
        ["cvrp_core.py", "CVRP instance generation, baseline heuristics, evaluation"],
        ["sample_efficient_search.py", "Sample-efficient search with duplicate filtering"],
        ["run_milestone.py", "Experiment entry point; outputs JSON results"],
        ["generate_report.py", "Auto-generate Markdown milestone report"],
        ["export_report_docx.py", "Export report as Word document (this file)"],
        ["outputs/milestone_results.json", "Experiment summary and best heuristic"],
        ["outputs/search_history.json", "Detailed search trajectory (all candidates + pruning info)"],
    ]
    add_table(doc, code_files, header=True)

    # 8. Reproducibility
    doc.add_paragraph()
    add_heading(doc, "8. Reproducibility", level=2)

    add_text(
        doc,
        "To reproduce results:\n"
        "1. cd milestone_cvrp\n"
        "2. python run_milestone.py\n"
        "3. python generate_report.py\n"
        "\n"
        "All random seeds are fixed (seed=2026), ensuring identical results across runs.",
        size=10,
    )

    # Save
    out_path = Path("milestone_report_detailed.docx")
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
