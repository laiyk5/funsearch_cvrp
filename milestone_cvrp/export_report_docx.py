from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt


def add_heading(doc: Document, text: str, size: int = 13) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)


def add_text(doc: Document, text: str, size: int = 11) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)


def main() -> None:
    out_dir = Path("outputs")
    results = json.loads((out_dir / "milestone_results.json").read_text(encoding="utf-8"))

    baseline = results["baseline"]
    method = results["sample_efficient_funsearch"]
    improvement = results["improvement"]

    doc = Document()
    add_heading(doc, "Milestone Report Draft: Sample-Efficient FunSearch for CVRP", size=14)

    add_heading(doc, "Group Members", size=12)
    add_text(doc, "QIN YUANCHENG (59061061)")
    add_text(doc, "QIN ZIHENG (59866035)")
    add_text(doc, "LAI YIKAI (59563061)")

    add_heading(doc, "1. Problem Description and Motivation", size=12)
    add_text(doc, "We study the Capacitated Vehicle Routing Problem (CVRP), where a fleet of vehicles with fixed capacity must serve customer demands with minimum total travel distance.")
    add_text(doc, "Our topic is sample-efficient FunSearch for CVRP, which treats heuristic design as a search problem over executable programs and aims to reduce expensive evaluations.")

    add_heading(doc, "2. Method/Approach Design", size=12)
    add_text(doc, "Baseline: Nearest Neighbor routing with capacity feasibility.")
    add_text(doc, "Sample-efficient prototype: weighted greedy heuristic + mutation search + duplicate signature filtering + two-stage evaluation (early pruning + full evaluation).")

    add_heading(doc, "3. Preliminary Results", size=12)
    add_text(doc, f"Nearest Neighbor: avg distance = {baseline['avg_distance']}, avg routes = {baseline['avg_num_routes']}")
    add_text(doc, f"Sample-efficient FunSearch: avg distance = {method['avg_distance']}, avg routes = {method['avg_num_routes']}")
    add_text(doc, f"Distance improvement: {improvement['distance_delta']} ({improvement['distance_improvement_percent']}%), evaluated unique candidates: {method['unique_candidates']}")

    add_heading(doc, "4. Submitted Code", size=12)
    add_text(doc, "cvrp_core.py, sample_efficient_search.py, run_milestone.py, outputs/milestone_results.json, outputs/search_history.json")

    add_heading(doc, "5. Next Steps", size=12)
    add_text(doc, "Integrate CVRPLib/Uchoa benchmarks, add stronger baselines (OR-Tools/LKH), and run larger ablation studies.")

    out_path = Path("milestone_report.docx")
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
